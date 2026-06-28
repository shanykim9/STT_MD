import streamlit as st
import subprocess
import os
import tempfile
import requests
import re
import fitz  # PyMuPDF
from dotenv import load_dotenv
import glob
from openai import OpenAI
import io
import datetime
import yfinance as yf
import pandas as pd
import json
import docx  # python-docx 패키지

# 마크다운 -> 워드 변환용 패키지
import markdown
from docx import Document
from htmldocx import HtmlToDocx

# LlamaParse 패키지 및 비동기 방어
from llama_parse import LlamaParse
import nest_asyncio
nest_asyncio.apply()

# Anthropic (Claude) 패키지
import anthropic

# .env 파일 로드
load_dotenv()

# ==========================================
# 탭 4(시뮬레이션)를 위한 전역 함수 및 설정
# ==========================================
STATE_FILE = "portfolio_state.json"

def load_state():
    """로컬 JSON 파일에서 기존 시뮬레이션 상태를 불러옵니다."""
    if os.path.exists(STATE_FILE):
        with open(STATE_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {
        "cash": 100_000_000, 
        "portfolio": {}, 
        "pending_orders": [], 
        "trade_logs": [], 
        "last_lecture_date": None
    }

def save_state(state):
    """현재 시뮬레이션 상태를 로컬 JSON 파일에 덮어씁니다."""
    with open(STATE_FILE, "w", encoding="utf-8") as f:
        json.dump(state, f, ensure_ascii=False, indent=4)

def parse_docx_views(file_bytes):
    """업로드된 워드 문서에서 강의일과 매매뷰 표를 자동으로 추출합니다."""
    doc = docx.Document(io.BytesIO(file_bytes))
    lecture_date = None
    views = {}

    # 1. 강의 날짜 추출 (예: '26년 6월 15일)
    date_pattern = re.compile(r"(\d{2})년\s*(\d{1,2})월\s*(\d{1,2})일")
    for para in doc.paragraphs:
        match = date_pattern.search(para.text)
        if match:
            yy, mm, dd = match.groups()
            lecture_date = datetime.datetime(2000 + int(yy), int(mm), int(dd))
            break
            
    if not lecture_date:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    match = date_pattern.search(cell.text)
                    if match:
                        yy, mm, dd = match.groups()
                        lecture_date = datetime.datetime(2000 + int(yy), int(mm), int(dd))
                        break
                if lecture_date: break
            if lecture_date: break

    # 2. 표에서 매매뷰 추출
    for table in doc.tables:
        for row in table.rows[1:]: # 헤더 건너뛰기
            cols = row.cells
            if len(cols) >= 2:
                stock_name = cols[0].text.strip()
                view_text = cols[1].text.strip()
                if stock_name and view_text:
                    views[stock_name] = view_text
    return lecture_date, views

def get_action_type(view_text):
    """AI가 작성한 다양한 매매뷰 텍스트를 정형화된 액션 룰로 매핑합니다."""
    vc = view_text.replace("✅", "").replace("⚠️", "").replace("🔴", "").replace(" ", "")
    if "적극분할매수" in vc: return "적극 분할 매수"
    if "눌림목분할매수" in vc: return "눌림목 분할 매수"
    if "분할매수" in vc: return "분할 매수"
    if "소량장기보유" in vc: return "소량 장기보유"
    
    if "보수적관망" in vc: return "보수적 관망"
    if "삐뽀삐뽀" in vc: return "삐뽀삐뽀"
    
    if any(x in vc for x in ["관망", "소량보유", "소량관찰", "시큰둥"]): 
        return "매도 1/3"
        
    return "보유/대기"

# ==========================================
# 웹 브라우저 UI 시작점
# ==========================================
st.set_page_config(page_title="AI 종합 데이터 변환 및 분석 툴", page_icon="🛠️", layout="wide")

st.title("🛠️ AI 종합 데이터 변환 및 분석 툴")
st.markdown("영상 음성 추출, 고품질 PDF 복원, Claude 심층 분석, 그리고 **주차별 누적이 가능한 자동 매매 시뮬레이션**을 통합 제공합니다.")

tab1, tab2, tab3, tab4 = st.tabs(["🎥 영상 -> TXT 변환", "📄 PDF -> MD 변환", "🧠 AI 강의 종합 분석", "📈 누적형 매매 시뮬레이션"])

# ==========================================
# 탭 1: 영상 음성 텍스트 변환 (OpenAI Whisper / 다중 파일 업로드)
# ==========================================
with tab1:
    st.header("🎥 영상 음성 -> 텍스트 일괄 변환 (OpenAI Whisper)")
    st.markdown("네트워크 오류가 없는 가장 안정적인 **OpenAI Whisper API**를 사용합니다. 여러 영상을 한 번에 업로드할 수 있습니다.")
    
    openai_api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not openai_api_key:
        openai_api_key = st.text_input("OpenAI API Key를 입력하세요:", type="password", key="vid_key").strip()
        
    uploaded_videos = st.file_uploader("영상 파일들을 업로드하세요:", type=["mp4", "mkv", "avi", "mov"], accept_multiple_files=True, key="vid_upload")

    if st.button("🚀 안정적인 영상 일괄 변환 시작", key="btn_vid"):
        if not openai_api_key:
            st.error("API 키가 필요합니다.")
        elif not uploaded_videos:
            st.error("영상 파일을 하나 이상 업로드해 주세요.")
        else:
            try:
                client = OpenAI(api_key=openai_api_key)
                total_combined_transcript = ""
                
                for v_idx, video_file in enumerate(uploaded_videos):
                    st.write(f"### 🎬 [{v_idx+1}/{len(uploaded_videos)}] '{video_file.name}' 처리 중...")
                    
                    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(video_file.name)[1]) as tmp_vid:
                        tmp_vid.write(video_file.read())
                        video_path = tmp_vid.name

                    duration_cmd = ["ffprobe", "-v", "error", "-show_entries", "format=duration", "-of", "default=noprint_wrappers=1:nokey=1", video_path]
                    total_duration = float(subprocess.run(duration_cmd, stdout=subprocess.PIPE, text=True).stdout.strip())

                    progress_text = st.empty()
                    progress_bar = st.progress(0.0)

                    with tempfile.TemporaryDirectory() as temp_dir:
                        chunk_pattern = os.path.join(temp_dir, "chunk_%03d.mp3")
                        
                        command = [
                            "ffmpeg", "-y", "-i", video_path, "-vn", 
                            "-acodec", "libmp3lame", "-ac", "1", "-ar", "16000", "-b:a", "32k",
                            "-f", "segment", "-segment_time", "1800", 
                            chunk_pattern
                        ]
                        process = subprocess.Popen(command, stderr=subprocess.PIPE, universal_newlines=True, encoding='utf-8')
                        
                        time_regex = re.compile(r"time=(\d{2}):(\d{2}):(\d{2}\.\d{2})")
                        for line in process.stderr:
                            match = time_regex.search(line)
                            if match:
                                h, m, s = match.groups()
                                current_time = int(h) * 3600 + int(m) * 60 + float(s)
                                percent = min(1.0, current_time / total_duration)
                                progress_bar.progress(percent)
                                progress_text.text(f"1단계: 오디오 30분 단위 분할 중... ({int(percent * 100)}%)")
                        
                        process.wait()
                        progress_bar.progress(1.0)
                        progress_text.text("1단계 완료: 오디오 분할 추출 완료")

                        chunks = sorted(glob.glob(os.path.join(temp_dir, "chunk_*.mp3")))
                        total_chunks = len(chunks)
                        
                        video_transcript = f"=== [{video_file.name}] 변환 결과 ===\n\n"
                        has_error = False

                        for idx, chunk_file in enumerate(chunks):
                            progress_text.text(f"2단계: AI 음성 인식 중... (총 {total_chunks}개 조각 중 {idx+1}번째 처리 중 🔄)")
                            progress_bar.progress(idx / total_chunks)
                            
                            try:
                                with open(chunk_file, "rb") as audio_file:
                                    transcript = client.audio.transcriptions.create(
                                        model="whisper-1",
                                        file=audio_file,
                                        response_format="text"
                                    )
                                video_transcript += transcript + "\n\n"
                            except Exception as e:
                                st.error(f"오류: {idx+1}번째 조각 처리 중 에러 발생: {str(e)}")
                                has_error = True
                                break

                        if not has_error:
                            progress_bar.progress(1.0)
                            progress_text.text(f"✅ '{video_file.name}' 변환 완료!")
                            total_combined_transcript += video_transcript + "\n\n========================================\n\n"
                            
                    os.remove(video_path)

                st.success("✅ 모든 영상의 텍스트 변환이 완료되었습니다!")
                st.text_area("통합 변환 결과", total_combined_transcript, height=300)
                st.download_button("📄 통합 텍스트 다운로드", total_combined_transcript, file_name="영상_통합_변환결과.txt", mime="text/plain")

            except Exception as e:
                st.error(f"오류 발생: {str(e)}")

# ==========================================
# 탭 2: PDF 마크다운 변환 (LlamaParse / 다중 파일 업로드)
# ==========================================
with tab2:
    st.header("📄 고품질 PDF -> 마크다운(.md) 일괄 변환")
    llama_api_key = os.getenv("LLAMA_CLOUD_API_KEY", "").strip()
    if not llama_api_key:
        llama_api_key = st.text_input("LlamaCloud API Key:", type="password", key="pdf_key").strip()

    uploaded_pdfs = st.file_uploader("변환할 PDF 파일들을 업로드하세요:", type=["pdf"], key="pdf_upload", accept_multiple_files=True)

    if st.button("🚀 고품질 PDF 일괄 변환 시작", key="btn_pdf"):
        if not llama_api_key or not uploaded_pdfs:
            st.error("API Key와 PDF 파일이 최소 1개 이상 필요합니다.")
        else:
            try:
                parser = LlamaParse(api_key=llama_api_key, result_type="markdown", verbose=False)
                total_combined_markdown = ""

                for p_idx, pdf_file in enumerate(uploaded_pdfs):
                    st.write(f"### 📄 [{p_idx+1}/{len(uploaded_pdfs)}] '{pdf_file.name}' 분석 중...")
                    
                    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
                    total_pages = len(doc)

                    progress_text = st.empty()
                    progress_bar = st.progress(0.0)
                    
                    pdf_markdown = f"# --- [{pdf_file.name}] 변환 결과 ---\n\n"

                    for i in range(total_pages):
                        current_page = i + 1
                        progress_text.text(f"LlamaParse 분석 중... ({current_page}/{total_pages} 페이지)")
                        
                        page_doc = fitz.open()
                        page_doc.insert_pdf(doc, from_page=i, to_page=i)
                        
                        tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
                        tmp_page_path = tmp_file.name
                        tmp_file.close() 
                        
                        page_doc.save(tmp_page_path)
                        page_doc.close()
                        
                        parsed_docs = parser.load_data(tmp_page_path)
                        if parsed_docs:
                            pdf_markdown += parsed_docs[0].text + "\n\n"
                        
                        try: os.remove(tmp_page_path)
                        except: pass 

                        progress_bar.progress(current_page / total_pages)

                    progress_text.text(f"✅ '{pdf_file.name}' 마크다운 변환 완료!")
                    doc.close()
                    
                    total_combined_markdown += pdf_markdown + "\n\n"

                st.success("✅ 모든 PDF 파일의 마크다운 변환이 완료되었습니다!")
                st.text_area("통합 변환 결과", total_combined_markdown, height=300)
                st.download_button("📄 통합 마크다운 다운로드", total_combined_markdown, file_name="PDF_통합_변환결과.md", mime="text/markdown")
                
            except Exception as e:
                st.error(f"오류 발생: {str(e)}")

# ==========================================
# 탭 3: AI 종합 분석 (3단계 세분할 / 워드 다운로드 기능 추가)
# ==========================================
with tab3:
    st.header("🧠 주식 강의 무제한 심층 분석 (다중 문서 및 워드 추출 지원)")
    st.markdown("여러 개의 교재와 녹취록을 업로드하면, AI가 모든 자료를 취합하여 분석한 후 **워드 문서(.docx)** 또는 마크다운 파일로 제공합니다.")

    anthropic_api_key = os.getenv("ANTHROPIC_API_KEY", "").strip()
    if not anthropic_api_key:
        anthropic_api_key = st.text_input("Anthropic API Key를 입력하세요:", type="password", key="claude_key").strip()

    # --- 1. 화면 내 지침(Rule) 편집 영역 ---
    default_rules = """[주식 강의 교차 분석 및 맞춤형 요약 시스템 Rule / Skill]

지금 주식 강의 교재 파일(들)과 녹취록 파일(들)을 업로드합니다. 
이 자료들을 모두 교차 검증하여, 아래의 5가지 지침과 출력 형식에 맞춰 핵심 내용을 빠짐없이 심층 정리해 주세요.

**[분석 및 출력 지침]**
1. 종합 핵심 브리핑 (유연한 핵심 포인트 도출 및 요약 서술):
- 교재(PDF/MD)의 목차 흐름을 기본 뼈대로 하되, 강사가 녹취록(TXT)에서 특별히 시간을 많이 할애하여 열변을 토하거나 중요하다고 강조한 내용을 최우선 순위로 삼아 이번 주 최우선 핵심 포인트를 도출해 주세요.
- 핵심 포인트 개수는 기본 3가지로 하되, 강사가 강조한 내용이 많을 경우 5가지 혹은 그 이상으로 유연하게 모두 정리해 주세요. (적다면 1~2가지도 좋습니다.)
- (중요) 단순하고 짧은 요약을 지양하고, 강사가 강조한 핵심 배경과 이유를 이해할 수 있도록 요약해서 서술해 주세요.

2. 교재 + 강사의 추가 인사이트 결합:
- 교재에 있는 딱딱한 전문 용어나 개념 중, 강사가 녹취록에서 초보자도 이해하기 쉽게 비유를 들었거나 부연 설명한 부분을 적극적으로 찾아내어 그 설명 방식을 그대로 반영해 요약해서 정리해 주세요.

3. [강사 특별 코멘트] 섹션 분리:
- 교재에는 명시되어 있지 않으나 강사가 강의 중에 추가로 언급한 시장의 최신 동향, 거시경제(매크로) 노이즈에 대한 경고, 특정 섹터에 대한 개인적인 견해 및 주의사항이 있다면 [강사 특별 코멘트]라는 제목 아래 별도로 눈에 띄게 정리해 주세요. 배경 설명을 생략하지 말고 요약해서 기재해 주세요.

4. 종목별 심층 뷰(View) 및 밸류에이션 정리 (선택과 집중):
- 교재에 언급된 모든 종목을 다루되, 강사가 가장 강조한 상위 5~7개(강조한 내용이 많을 경우 5~7가지 이상으로 유연하게 모두 정리해 주세요. 적다면 1~2가지도 좋습니다.)
 핵심 종목만 아래 3가지 요소(A, B, C)를 포함하여 깊이 있게 작성해 주세요.
  A. 음성(TXT)에서 추가로 강조된 해당 종목의 핵심 모멘텀 및 이슈 (생략된 경우 해당 사실 기재)
  B. 해당 종목의 PER 수치 (교재와 녹취록에 언급된 수치 모두 기재, 없으면 '언급 없음' 기재)
  C. 담쌤(강사)의 명확한 매매 뷰 (예: 적극 분할 매수, 눌림목 대기, 관망 등)와 판단한 구체적인 이유 서술.
- 나머지 단순 언급되거나 패스한 종목들은 글자 수(토큰) 확보를 위해 서술형이 아닌 '요약 표(Table)' 형식(종목명, PER, 매매 뷰 3가지 컬럼)으로 압축해서 한 번에 정리해 주세요.

5. 실전 적용 포인트:
- 종합된 분석 내용을 바탕으로, 투자자가 이번 주 실전 투자나 추가 공부 시 가장 집중해야 할 구체적인 전략이나 결론을 2~3가지 요약해 주세요.
- 담쌤(강사)의 강의한 알짜를 맨 먼저 기재하고, 각 종목별 매매뷰를 간략히 핵심만 요약해서 표로 정리해주세요. (예: '26년 6월 1일(강의일), 삼성전자 적극매수, SK하이닉스 적극매수, 삼성전기 조정기다림 등')"""

    st.subheader("⚙️ AI 분석 지침 (프롬프트)")
    user_rules = st.text_area("이 지침을 바탕으로 AI가 모든 문서를 교차 검증합니다:", value=default_rules, height=350)
    st.markdown("---")

    # --- 2. 데이터 업로드 영역 (다중 업로드 허용) ---
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("📁 1. 교재 / 강의자료")
        material_files = st.file_uploader("교재 업로드 (다중 선택 가능, MD 또는 PDF)", type=["md", "pdf"], accept_multiple_files=True)
    with col2:
        st.subheader("📁 2. 강사 녹취록")
        transcript_files = st.file_uploader("녹취록 업로드 (다중 선택 가능, TXT)", type=["txt"], accept_multiple_files=True)

    if st.button("🚀 취합된 자료로 3단계 세분할 분석 시작", type="primary"):
        if not anthropic_api_key:
            st.error("Anthropic API Key가 필요합니다.")
        elif not material_files or not transcript_files:
            st.error("교재 파일과 녹취록 파일을 각각 하나 이상 업로드해 주세요.")
        else:
            try:
                with st.spinner("모든 파일의 텍스트를 하나로 취합하는 중입니다..."):
                    material_text = ""
                    for m_file in material_files:
                        material_text += f"\n\n--- [교재 자료: {m_file.name}] ---\n\n"
                        if m_file.name.endswith(".pdf"):
                            doc = fitz.open(stream=m_file.read(), filetype="pdf")
                            for page in doc:
                                material_text += page.get_text() + "\n"
                            doc.close()
                        else:
                            material_text += m_file.read().decode("utf-8")
                    
                    transcript_text = ""
                    for t_file in transcript_files:
                        transcript_text += f"\n\n--- [강사 녹취록: {t_file.name}] ---\n\n"
                        transcript_text += t_file.read().decode("utf-8")

                client = anthropic.Anthropic(api_key=anthropic_api_key)
                
                st.markdown("### 📊 분석 결과 리포트")
                st.markdown("---")
                part1_container = st.empty()
                part2_container = st.empty()
                part3_container = st.empty()
                
                # ==========================================
                # [1단계 호출] 지침 1, 2, 3 처리
                # ==========================================
                with st.spinner("🔄 파트 1 분석 중: 취합된 모든 자료를 바탕으로 지침 1~3번을 작성 중입니다..."):
                    system_prompt_1 = "당신은 세계 최고의 주식 강의 분석가입니다. 사용자가 제공하는 <instructions> 안의 지침 중 [1번, 2번, 3번]에 해당하는 내용만 완벽하게 작성하십시오. 4번과 5번은 절대 출력하지 마십시오."
                    
                    user_message_1 = f"""
<instructions>
{user_rules}

[중요 출력 규칙]
- 답변 생성 시 반드시 지침에 명시된 1, 2, 3번 목차 번호와 제목을 그대로 사용하십시오.
- 4번과 5번은 이 단계에서 절대로 출력하지 마십시오.
</instructions>

<material>
{material_text}
</material>

<transcript>
{transcript_text}
</transcript>
"""
                    response_1 = client.messages.create(
                        model="claude-sonnet-4-6", 
                        max_tokens=8192,
                        temperature=0.3,
                        system=system_prompt_1,
                        messages=[{"role": "user", "content": user_message_1}]
                    )
                    part1_result = response_1.content[0].text
                    part1_container.markdown(part1_result)

                # ==========================================
                # [2단계 호출] 지침 4 처리
                # ==========================================
                with st.spinner("🔄 파트 2 분석 중: 지침 4번(종목별 심층 뷰)을 집중 분석 중입니다..."):
                    system_prompt_2 = "당신은 세계 최고의 주식 강의 분석가입니다. 오직 <instructions> 안의 지침 중 [4번]에 해당하는 내용만 집중적으로 작성하십시오."
                    
                    user_message_2 = f"""
<instructions>
{user_rules}

[중요 출력 규칙]
- 답변 생성 시 반드시 지침에 명시된 4번 목차 번호와 제목을 그대로 사용하십시오.
- 오직 4번 지침만 수행하십시오. 1, 2, 3번이나 5번 내용을 출력하지 마십시오.
</instructions>

<material>
{material_text}
</material>

<transcript>
{transcript_text}
</transcript>
"""
                    response_2 = client.messages.create(
                        model="claude-sonnet-4-6", 
                        max_tokens=8192,
                        temperature=0.3,
                        system=system_prompt_2,
                        messages=[{"role": "user", "content": user_message_2}]
                    )
                    part2_result = response_2.content[0].text
                    part2_container.markdown(part2_result)

                # ==========================================
                # [3단계 호출] 지침 5 처리
                # ==========================================
                with st.spinner("🔄 파트 3 분석 중: 지침 5번(실전 적용 포인트 요약)을 마무리 중입니다..."):
                    system_prompt_3 = "당신은 세계 최고의 주식 강의 분석가입니다. 오직 <instructions> 안의 지침 중 [5번]에 해당하는 내용만 작성하십시오."
                    
                    user_message_3 = f"""
앞서 당신이 작성한 1~4번의 분석 결과와 제공된 원본 데이터를 바탕으로, 마지막 [5번] 지침을 완벽하게 수행하십시오.

<instructions>
{user_rules}

[중요 출력 규칙]
- 답변 생성 시 반드시 지침에 명시된 5번 목차 번호와 제목을 그대로 사용하십시오.
- 오직 5번 지침에 대해서만 답변하십시오. 
</instructions>

<material>
{material_text}
</material>

<transcript>
{transcript_text}
</transcript>

<previous_analysis>
{part1_result}
{part2_result}
</previous_analysis>
"""
                    response_3 = client.messages.create(
                        model="claude-sonnet-4-6", 
                        max_tokens=8192,
                        temperature=0.3,
                        system=system_prompt_3,
                        messages=[{"role": "user", "content": user_message_3}]
                    )
                    part3_result = response_3.content[0].text
                    part3_container.markdown(part3_result)

                st.success("✅ 여러 개의 자료를 바탕으로 한 3단계 심층 분석이 완벽하게 완료되었습니다!")
                
                # 파일 다운로드 처리 로직 (Markdown & Word)
                full_report = part1_result + "\n\n" + part2_result + "\n\n" + part3_result
                
                html_content = markdown.markdown(full_report, extensions=['tables'])
                doc = Document()
                new_parser = HtmlToDocx()
                new_parser.add_html_to_document(html_content, doc)
                
                doc_io = io.BytesIO()
                doc.save(doc_io)
                docx_data = doc_io.getvalue()
                
                col_btn1, col_btn2 = st.columns(2)
                with col_btn1:
                    st.download_button(
                        label="📥 마크다운 리포트 다운로드 (.md)",
                        data=full_report,
                        file_name="주식강의_다중문서_심층분석_리포트.md",
                        mime="text/markdown",
                        use_container_width=True
                    )
                with col_btn2:
                    st.download_button(
                        label="📝 워드 문서 리포트 다운로드 (.docx)",
                        data=docx_data,
                        file_name="주식강의_다중문서_심층분석_리포트.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                        use_container_width=True
                    )

            except Exception as e:
                st.error(f"분석 중 오류가 발생했습니다: {str(e)}")

# ==========================================
# 탭 4: 주식 매매 시뮬레이션 (누적 저장 및 자동 파싱 지원)
# ==========================================
with tab4:
    st.header("📈 주차별 누적 매매 시뮬레이션 및 포트폴리오 관리")
    st.markdown("매주 생성된 '종목별 매매뷰' 워드 문서를 업로드하면, **기존의 현금 잔고와 주식 보유량을 기억(Stateful)**하고 이어서 매매를 진행합니다.")
    
    TICKER_MAP = {
        "SK하이닉스": "000660.KS", "삼성전자": "005930.KS", "마벨테크놀로지": "MRVL",
        "LIG D&A": "079550.KS", "GE버노바": "GEV", "현대모비스": "012330.KS",
        "삼양식품": "003230.KS", "엔비디아": "NVDA", "TSMC": "TSM",
        "SK텔레콤": "017670.KS", "삼성생명": "032830.KS", "삼성증권": "016360.KS",
        "샌디스크": "WDC", "삼성전기": "009150.KS", "한화에어로스페이스": "012450.KS",
        "현대차": "005380.KS", "메타": "META", "알파벳": "GOOGL",
        "아마존": "AMZN", "테슬라": "TSLA", "스페이스X": None, 
        "브로드컴": "AVGO", "마이크로소프트": "MSFT", "팔란티어": "PLTR",
        "이튼": "ETN", "센트러스에너지": "LEU"
    }

    col_btn1, col_btn2 = st.columns([2, 1])
    with col_btn2:
        if st.button("🔄 잔고 및 시뮬레이션 초기화 (1억 원부터 다시)"):
            if os.path.exists(STATE_FILE):
                os.remove(STATE_FILE)
            st.success("✅ 시뮬레이션 데이터가 성공적으로 초기화되었습니다.")
            st.rerun()

    uploaded_view_doc = st.file_uploader("이번 주 '종목별 매매뷰' 워드 파일(.docx)을 업로드하세요:", type=["docx"], key="view_doc")

    if st.button("🚀 누적 시뮬레이션 실행 및 결과 덮어쓰기", type="primary"):
        if not uploaded_view_doc:
            st.error("매매뷰 워드 파일을 업로드해 주세요.")
        else:
            with st.spinner("문서 해독 및 과거 데이터 백테스팅을 진행 중입니다..."):
                try:
                    lecture_date, parsed_views = parse_docx_views(uploaded_view_doc.read())
                    if not lecture_date:
                        st.error("워드 문서에서 강의일(예: '26년 6월 15일')을 찾을 수 없습니다.")
                        st.stop()

                    state = load_state()
                    last_date_str = state["last_lecture_date"]
                    
                    if last_date_str:
                        last_date = datetime.datetime.strptime(last_date_str, "%Y-%m-%d")
                        if lecture_date <= last_date:
                            st.warning(f"⚠️ 업로드하신 문서의 강의일({lecture_date.strftime('%Y-%m-%d')})이 이미 처리된 시뮬레이션({last_date.strftime('%Y-%m-%d')})과 같거나 과거입니다. 새로운 주차의 문서를 업로드해 주세요.")
                            st.stop()
                    else:
                        last_date = lecture_date - datetime.timedelta(days=365)
                    
                    today = datetime.datetime.today()
                    start_date = last_date

                    prices = {}
                    for name, ticker in TICKER_MAP.items():
                        if ticker:
                            try:
                                hist = yf.Ticker(ticker).history(start=start_date, end=today)
                                if not hist.empty:
                                    hist.index = hist.index.tz_localize(None)
                                    prices[ticker] = hist
                            except:
                                pass
                    
                    usd_krw = yf.Ticker("USDKRW=X").history(start=start_date, end=today)
                    if not usd_krw.empty:
                        usd_krw.index = usd_krw.index.tz_localize(None)

                    # 1. 과거 대기 주문(Pending Orders) 트래킹 (지난주 ~ 이번주 강의일)
                    unexecuted_orders = []
                    for order in state["pending_orders"]:
                        ticker = order['ticker']
                        if ticker not in prices: continue
                        df = prices[ticker]
                        
                        order_date = datetime.datetime.strptime(order['created_date'], "%Y-%m-%d")
                        check_df = df[(df.index >= order_date) & (df.index <= lecture_date)]
                        
                        executed = False
                        for date, row in check_df.iterrows():
                            if row['Low'] <= order['target']:
                                rate = usd_krw['Close'].reindex(df.index, method='ffill').loc[date] if order['is_us'] and not usd_krw.empty else 1
                                cost_krw = order['shares'] * order['target'] * rate
                                if state["cash"] >= cost_krw:
                                    state["cash"] -= cost_krw
                                    if order['name'] not in state["portfolio"]: 
                                        state["portfolio"][order['name']] = {'shares': 0, 'total_cost': 0}
                                    state["portfolio"][order['name']]['shares'] += order['shares']
                                    state["portfolio"][order['name']]['total_cost'] += cost_krw
                                    state["trade_logs"].append(f"⚡ [대기 체결] {date.strftime('%Y-%m-%d')} | {order['name']} | 목표가 {order['target']:.2f} 터치 | {order['shares']}주 편입")
                                    executed = True
                                    break
                        if not executed:
                            unexecuted_orders.append(order)
                    
                    state["pending_orders"] = unexecuted_orders

                    # 2. 신규 뷰 적용 (이번주 강의일)
                    state["trade_logs"].append(f"\n--- 📅 [{lecture_date.strftime('%Y-%m-%d')} 강의 시그널 적용] ---")
                    
                    for name, view_text in parsed_views.items():
                        ticker = TICKER_MAP.get(name)
                        if not ticker or ticker not in prices: continue
                        
                        action = get_action_type(view_text)
                        df = prices[ticker]
                        
                        before_lecture = df[df.index < lecture_date]
                        if before_lecture.empty: continue
                        prev_close = before_lecture['Close'].iloc[-1]
                        
                        on_after_lecture = df[df.index >= lecture_date]
                        if on_after_lecture.empty: continue
                        lecture_day = on_after_lecture.iloc[0]
                        lecture_open = lecture_day['Open']
                        actual_date = lecture_day.name
                        
                        is_us = not (ticker.endswith(".KS") or ticker.endswith(".KQ"))
                        rate = usd_krw['Close'].reindex(df.index, method='ffill').loc[actual_date] if is_us and not usd_krw.empty else 1
                        
                        # [매수 로직]
                        if action == "적극 분할 매수":
                            price_krw = lecture_open * rate
                            shares = 1 if price_krw > 3_000_000 else int(3_000_000 // price_krw)
                            cost_krw = shares * price_krw
                            if shares > 0 and state["cash"] >= cost_krw:
                                state["cash"] -= cost_krw
                                if name not in state["portfolio"]: state["portfolio"][name] = {'shares': 0, 'total_cost': 0}
                                state["portfolio"][name]['shares'] += shares
                                state["portfolio"][name]['total_cost'] += cost_krw
                                state["trade_logs"].append(f"✅ [시가 매수] {name} | {shares}주 | 매수금: 약 {int(cost_krw):,}원")
                                
                        elif action == "분할 매수":
                            target_price = prev_close * 0.98
                            target_krw = target_price * rate
                            shares = 1 if target_krw > 2_000_000 else int(2_000_000 // target_krw)
                            if shares > 0:
                                state["pending_orders"].append({'name': name, 'ticker': ticker, 'target': target_price, 'shares': shares, 'is_us': is_us, 'tag': '분할 매수 대기', 'created_date': actual_date.strftime("%Y-%m-%d")})
                                state["trade_logs"].append(f"⏳ [주문 예약] {name} | 목표가: {target_price:.2f} | {shares}주 대기")
                                
                        elif action == "눌림목 분할 매수":
                            target_price = prev_close * 0.95
                            target_krw = target_price * rate
                            shares = 1 if target_krw > 3_000_000 else int(3_000_000 // target_krw)
                            if shares > 0:
                                state["pending_orders"].append({'name': name, 'ticker': ticker, 'target': target_price, 'shares': shares, 'is_us': is_us, 'tag': '눌림목 대기', 'created_date': actual_date.strftime("%Y-%m-%d")})
                                state["trade_logs"].append(f"⏳ [주문 예약] {name} | 목표가: {target_price:.2f} | {shares}주 대기")
                                
                        elif action == "소량 장기보유":
                            price_krw = lecture_open * rate
                            shares = 1 if price_krw > 500_000 else int(500_000 // price_krw)
                            cost_krw = shares * price_krw
                            if shares > 0 and state["cash"] >= cost_krw:
                                state["cash"] -= cost_krw
                                if name not in state["portfolio"]: state["portfolio"][name] = {'shares': 0, 'total_cost': 0}
                                state["portfolio"][name]['shares'] += shares
                                state["portfolio"][name]['total_cost'] += cost_krw
                                state["trade_logs"].append(f"✅ [1차 매수] {name} | {shares}주")
                                t2, t3 = lecture_open * 0.97, lecture_open * 0.97 * 0.97
                                s2 = 1 if (t2 * rate) > 500_000 else int(500_000 // (t2 * rate))
                                s3 = 1 if (t3 * rate) > 500_000 else int(500_000 // (t3 * rate))
                                state["pending_orders"].append({'name': name, 'ticker': ticker, 'target': t2, 'shares': s2, 'is_us': is_us, 'tag': '2차 매수', 'created_date': actual_date.strftime("%Y-%m-%d")})
                                state["pending_orders"].append({'name': name, 'ticker': ticker, 'target': t3, 'shares': s3, 'is_us': is_us, 'tag': '3차 매수', 'created_date': actual_date.strftime("%Y-%m-%d")})

                        # [매도 로직] - 현재 보유 중일 때만 작동
                        elif action in ["매도 1/3", "보수적 관망", "삐뽀삐뽀"]:
                            current_holding = state["portfolio"].get(name, {}).get("shares", 0)
                            if current_holding > 0:
                                sell_shares = 0
                                if action == "매도 1/3": sell_shares = max(1, current_holding // 3)
                                elif action == "보수적 관망": sell_shares = max(1, current_holding // 2)
                                elif action == "삐뽀삐뽀": sell_shares = current_holding
                                
                                if sell_shares > 0:
                                    avg_cost = state["portfolio"][name]["total_cost"] / current_holding
                                    sell_amount_krw = sell_shares * lecture_open * rate
                                    realized_pnl = sell_amount_krw - (sell_shares * avg_cost)
                                    
                                    state["cash"] += sell_amount_krw
                                    state["portfolio"][name]["shares"] -= sell_shares
                                    state["portfolio"][name]["total_cost"] -= (sell_shares * avg_cost)
                                    
                                    state["trade_logs"].append(f"🚨 [매도 체결] {name} | {action} | {sell_shares}주 처분 | 회수: 약 {int(sell_amount_krw):,}원 (실현손익: {int(realized_pnl):,}원)")
                                    
                                    if state["portfolio"][name]["shares"] == 0:
                                        del state["portfolio"][name]

                    # 최신 상태 JSON 저장
                    state["last_lecture_date"] = lecture_date.strftime("%Y-%m-%d")
                    save_state(state)

                    # 3. 실시간 현재가 기준 포트폴리오 평가
                    evaluation_results = []
                    total_portfolio_value = 0

                    for name, data in state["portfolio"].items():
                        if data['shares'] > 0:
                            ticker = TICKER_MAP[name]
                            is_us = not (ticker.endswith(".KS") or ticker.endswith(".KQ"))
                            
                            current_local_price = prices[ticker]['Close'].iloc[-1]
                            current_rate = usd_krw['Close'].iloc[-1] if is_us and not usd_krw.empty else 1
                            
                            current_price_krw = current_local_price * current_rate
                            avg_cost_krw = data['total_cost'] / data['shares']
                            current_value_krw = current_price_krw * data['shares']
                            
                            pnl_amount = current_value_krw - data['total_cost']
                            pnl_pct = (pnl_amount / data['total_cost']) * 100
                            
                            status = "🔴 수익" if pnl_amount > 0 else ("🔵 손실" if pnl_amount < 0 else "⚪ 보합")
                            total_portfolio_value += current_value_krw

                            evaluation_results.append({
                                "상태": status,
                                "종목명": name,
                                "보유수량": int(data['shares']),
                                "평균단가(원)": f"{int(avg_cost_krw):,}",
                                "현재가(원)": f"{int(current_price_krw):,}",
                                "총매수금액(원)": f"{int(data['total_cost']):,}",
                                "현재평가금액(원)": f"{int(current_value_krw):,}",
                                "손익금액(원)": f"{int(pnl_amount):,}",
                                "손익률(%)": f"{round(pnl_pct, 2)}%" 
                            })

                    total_assets = state["cash"] + total_portfolio_value
                    total_profit = total_assets - 100_000_000
                    total_profit_pct = (total_profit / 100_000_000) * 100

                    # UI 출력
                    st.success("✅ 주차별 문서 반영 및 포트폴리오 누적 평가가 완료되었습니다!")
                    
                    col1, col2, col3, col4 = st.columns(4)
                    with col1: st.metric("💰 초기 자본금", "100,000,000 원")
                    with col2: st.metric("💵 현재 잔여 현금", f"{int(state['cash']):,} 원")
                    with col3: st.metric("📈 주식 평가 금액", f"{int(total_portfolio_value):,} 원")
                    with col4: st.metric("🏆 총 자산 (손익률)", f"{int(total_assets):,} 원", f"{total_profit_pct:+.2f}%")

                    st.subheader("📊 현재 포트폴리오 수익 상태")
                    if evaluation_results:
                        df_eval = pd.DataFrame(evaluation_results)
                        st.dataframe(df_eval, use_container_width=True)
                    else:
                        st.info("현재 보유 중인 주식이 없습니다.")

                    st.subheader("📝 누적 상세 거래 로그")
                    if not state["trade_logs"]:
                        st.info("아직 거래 내역이 없습니다.")
                    for log in reversed(state["trade_logs"]): 
                        st.text(log)

                except Exception as e:
                    st.error(f"시뮬레이션 중 오류가 발생했습니다. (사유: {str(e)})")